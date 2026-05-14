"""
PDF (and DOCX) text extraction for the budget pipeline.

Responsibilities:
  - Extract text from PDF files, page by page (via PyMuPDF direct text first,
    pytesseract OCR fallback for scanned pages).
  - Extract text from .docx files (Australian Appropriation Acts).
  - Cache raw extracted text per file to avoid re-OCR on repeated runs.
  - Return a list of PageText objects: [(page_num, text), ...]

The cache is shared with budget/ (same full_text directory) where files already
exist, so PDFs already OCR'd won't be re-processed.
"""

from __future__ import annotations

import gzip
import hashlib
import json
import logging
import shutil
import re
import struct
from dataclasses import dataclass
from pathlib import Path
from typing import Optional

logger = logging.getLogger(__name__)

# ---------------------------------------------------------------------------
# Data types
# ---------------------------------------------------------------------------

@dataclass
class PageText:
    page_num: int         # 1-based
    text: str
    method: str           # "direct" | "ocr" | "docx"
    char_count: int = 0

    def __post_init__(self):
        if not self.char_count:
            self.char_count = len(self.text)


# ---------------------------------------------------------------------------
# Cache helpers
# ---------------------------------------------------------------------------

def _file_id(path: Path) -> str:
    """Return a stable short hash of the file path (not content) for cache keys."""
    return hashlib.md5(str(path).encode()).hexdigest()[:12]


def _cache_path(cache_dir: Path, file_id: str) -> Path:
    return cache_dir / f"{file_id}.json.gz"


def _load_cache(cache_dir: Path, file_id: str) -> Optional[list[dict]]:
    p = _cache_path(cache_dir, file_id)
    if not p.exists():
        return None
    try:
        with gzip.open(p, "rt", encoding="utf-8") as f:
            return json.load(f)
    except Exception as e:
        logger.warning(f"Cache read error for {file_id}: {e}")
        return None


def _save_cache(cache_dir: Path, file_id: str, pages: list[PageText]) -> None:
    cache_dir.mkdir(parents=True, exist_ok=True)
    p = _cache_path(cache_dir, file_id)
    data = [{"page_num": pg.page_num, "text": pg.text, "method": pg.method} for pg in pages]
    try:
        with gzip.open(p, "wt", encoding="utf-8") as f:
            json.dump(data, f, ensure_ascii=False)
    except Exception as e:
        logger.warning(f"Cache write error for {file_id}: {e}")


def _deserialize_pages(data: list[dict]) -> list[PageText]:
    return [PageText(page_num=d["page_num"], text=d["text"], method=d["method"]) for d in data]


# ---------------------------------------------------------------------------
# PDF extraction (PyMuPDF + tesseract fallback)
# ---------------------------------------------------------------------------

_MIN_DIRECT_TEXT_CHARS = 80   # threshold below which we run OCR
_MIN_ALNUM_RATIO = 0.20
# Italian budget PDFs (2001-2016) have landscape-rotated financial tables whose
# text layer is extracted character-by-character in vertical reading order,
# producing garbled output like "DIRE Z IONE" instead of "DIREZIONE".
# This manifests as an unusually high proportion of single-character tokens.
_MAX_SINGLE_CHAR_TOKEN_RATIO = 0.28


def _text_is_usable(text: str) -> bool:
    if len(text) < _MIN_DIRECT_TEXT_CHARS:
        return False
    alnum = sum(c.isalnum() for c in text)
    if (alnum / max(len(text), 1)) < _MIN_ALNUM_RATIO:
        return False
    # Detect garbled vertical/rotated text: too many 1-char tokens signals
    # that the PDF text layer was read column-by-column instead of row-by-row.
    tokens = text.split()
    if len(tokens) >= 20:
        single_char_ratio = sum(1 for t in tokens if len(t) == 1) / len(tokens)
        if single_char_ratio > _MAX_SINGLE_CHAR_TOKEN_RATIO:
            return False
    return True


def _extract_pdf(
    path: Path,
    ocr_zoom: float = 2.0,
    ocr_langs: str = "eng",
    force_ocr: bool = False,
) -> list[PageText]:
    """Extract text from a PDF, page by page.

    Falls back to OCR on sparse pages or pages with garbled rotated text.
    Set force_ocr=True to skip the direct-text layer entirely (useful for
    countries whose PDFs always have unreliable embedded text layers, e.g.
    Italy 2001-2016 landscape budget tables).
    """
    try:
        import fitz  # PyMuPDF
    except ImportError:
        raise ImportError("PyMuPDF (fitz) required: pip install pymupdf")

    pages: list[PageText] = []
    if hasattr(fitz, "TOOLS"):
        try:
            fitz.TOOLS.mupdf_display_errors(False)
            fitz.TOOLS.mupdf_display_warnings(False)
            fitz.TOOLS.reset_mupdf_warnings()
        except Exception:
            pass

    doc = fitz.open(str(path))
    try:
        for page_index in range(len(doc)):
            page_obj = doc[page_index]
            page_num = page_index + 1
            try:
                direct_text = page_obj.get_text("text")
            except Exception as e:
                logger.debug("Direct text extraction failed on %s page %s: %s", path.name, page_num, e)
                direct_text = ""

            use_direct = (not force_ocr) and _text_is_usable(direct_text)

            if use_direct:
                pages.append(PageText(page_num=page_num, text=direct_text.strip(), method="direct"))
            else:
                # Try OCR — use auto-orientation mode (psm 1) so Tesseract detects
                # and corrects landscape / rotated pages automatically.
                ocr_text = _ocr_page(
                    page_obj,
                    zoom=ocr_zoom,
                    langs=ocr_langs,
                    # psm 1 = auto with OSD (detects 90°/180°/270° rotation).
                    # Falls back to psm 6 if OSD fails.
                    use_osd=True,
                )
                if ocr_text:
                    pages.append(PageText(page_num=page_num, text=ocr_text.strip(), method="ocr"))
                else:
                    # Last resort: keep whatever direct text we have
                    pages.append(PageText(page_num=page_num, text=direct_text.strip(), method="direct"))
    finally:
        doc.close()

    if hasattr(fitz, "TOOLS"):
        try:
            suppressed = fitz.TOOLS.mupdf_warnings()
            if suppressed:
                logger.debug("MuPDF suppressed warnings for %s: %s", path.name, suppressed)
        except Exception:
            pass
    return pages


def _ocr_page(page_obj, zoom: float, langs: str, use_osd: bool = False) -> str:
    """OCR a single PDF page using pytesseract.

    use_osd=True uses --psm 1 (auto with orientation and script detection),
    which lets Tesseract auto-correct landscape / rotated pages.  Falls back
    to --psm 6 if the OSD pass raises an error (e.g. missing osd.traineddata).
    """
    try:
        import pytesseract
        from PIL import Image
        import fitz
    except ImportError:
        logger.warning("pytesseract / PIL not available — skipping OCR")
        return ""

    try:
        tesseract_cmd = shutil.which("tesseract")
        if not tesseract_cmd:
            for candidate in ("/opt/homebrew/bin/tesseract", "/usr/local/bin/tesseract"):
                if Path(candidate).exists():
                    tesseract_cmd = candidate
                    break
        if tesseract_cmd:
            pytesseract.pytesseract.tesseract_cmd = tesseract_cmd

        mat = fitz.Matrix(zoom, zoom)
        pix = page_obj.get_pixmap(matrix=mat, alpha=False)
        img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)

        if use_osd:
            try:
                # psm 1: automatic page segmentation with OSD — handles rotation
                result = pytesseract.image_to_string(img, lang=langs, config="--psm 1")
                if result.strip():
                    return result
            except Exception as osd_err:
                logger.debug(f"OSD OCR failed ({osd_err}), falling back to psm 6")

        return pytesseract.image_to_string(img, lang=langs, config="--psm 6")
    except Exception as e:
        logger.debug(f"OCR failed: {e}")
        return ""


# ---------------------------------------------------------------------------
# DOCX extraction (Australian Appropriation Acts)
# ---------------------------------------------------------------------------

def _extract_docx(path: Path) -> list[PageText]:
    """Extract text from a .docx file. Returns one PageText per paragraph batch."""
    try:
        import docx
    except ImportError:
        raise ImportError("python-docx required: pip install python-docx")

    doc = docx.Document(str(path))
    all_text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())

    # Also extract tables
    table_rows: list[str] = []
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            table_rows.append("\t".join(cells))
    if table_rows:
        all_text += "\n\n--- TABLES ---\n" + "\n".join(table_rows)

    if not all_text.strip():
        try:
            import subprocess
            proc = subprocess.run(
                ["textutil", "-convert", "txt", "-stdout", str(path)],
                capture_output=True,
                text=True,
                check=True,
            )
            all_text = proc.stdout
        except Exception:
            pass

    # Split into ~page-sized chunks (3000 chars each) for chunked processing
    chunk_size = 3000
    chunks = [all_text[i : i + chunk_size] for i in range(0, len(all_text), chunk_size)]
    pages: list[PageText] = []
    for i, chunk in enumerate(chunks):
        pages.append(PageText(page_num=i + 1, text=chunk.strip(), method="docx"))
    return pages


def _extract_doc(path: Path) -> list[PageText]:
    """Extract text from a legacy .doc file using platform tools."""
    import subprocess

    text = ""
    last_error: Exception | None = None

    # macOS textutil handles legacy .doc reasonably well and is available in this environment.
    try:
        proc = subprocess.run(
            ["textutil", "-convert", "txt", "-stdout", str(path)],
            capture_output=True,
            text=True,
            check=True,
        )
        text = proc.stdout
    except Exception as e:
        last_error = e

    if not text.strip():
        if last_error:
            raise RuntimeError(f"Could not extract legacy .doc file {path.name}: {last_error}")
        raise RuntimeError(f"Could not extract legacy .doc file {path.name}")

    chunk_size = 3000
    chunks = [text[i : i + chunk_size] for i in range(0, len(text), chunk_size)]
    pages: list[PageText] = []
    for i, chunk in enumerate(chunks):
        pages.append(PageText(page_num=i + 1, text=chunk.strip(), method="docx"))
    return pages


# ---------------------------------------------------------------------------
# Public interface
# ---------------------------------------------------------------------------

def extract_pages(
    path: Path,
    cache_dir: Optional[Path] = None,
    force_reextract: bool = False,
    ocr_zoom: float = 2.0,
    ocr_langs: str = "eng",
    force_ocr: bool = False,
) -> list[PageText]:
    """
    Extract text from a PDF or DOCX file, using cache if available.

    Args:
        path:             Absolute path to the source file.
        cache_dir:        Directory for the text cache. If None, no caching.
        force_reextract:  If True, ignore cache and re-extract.
        ocr_zoom:         Pixel zoom factor for OCR (higher = better quality, slower).
        ocr_langs:        Tesseract language codes, e.g. "eng", "eng+fra+dan".
        force_ocr:        If True, skip the direct text layer and always OCR (useful for
                          countries whose PDFs have unreliable embedded text layers).

    Returns:
        List of PageText objects, one per page (1-indexed).
    """
    suffix = path.suffix.lower()

    # Try cache first
    if cache_dir and not force_reextract:
        file_id = _file_id(path)
        cached = _load_cache(cache_dir, file_id)
        if cached is not None:
            logger.debug(f"Cache hit: {path.name}")
            return _deserialize_pages(cached)

    if path.stat().st_size == 0:
        logger.warning(f"Skipping empty file (0 bytes): {path.name}")
        return []

    logger.info(f"Extracting text from: {path.name}")

    if suffix == ".pdf":
        pages = _extract_pdf(path, ocr_zoom=ocr_zoom, ocr_langs=ocr_langs, force_ocr=force_ocr)
    elif suffix == ".docx":
        pages = _extract_docx(path)
    elif suffix == ".doc":
        pages = _extract_doc(path)
    else:
        raise ValueError(f"Unsupported file format: {suffix}")

    # Save to cache
    if cache_dir:
        file_id = _file_id(path)
        _save_cache(cache_dir, file_id, pages)

    return pages


def pages_to_text(pages: list[PageText]) -> str:
    """Concatenate pages into a single string with page separators."""
    parts: list[str] = []
    for pg in pages:
        parts.append(f"[PAGE {pg.page_num}]\n{pg.text}")
    return "\n\n".join(parts)


def chunk_pages(
    pages: list[PageText],
    chunk_size: int = 15_000,
    overlap: int = 300,
    max_pages: int = 10,
) -> list[tuple[list[PageText], str]]:
    """
    Split pages into chunks suitable for a single LLM call.

    Returns:
        List of (page_list, combined_text) tuples.
        Each chunk has at most max_pages pages and at most chunk_size chars.
    """
    chunks: list[tuple[list[PageText], str]] = []
    current_pages: list[PageText] = []
    current_chars = 0

    for pg in pages:
        pg_chars = len(pg.text)

        # Start a new chunk if adding this page would exceed limits
        if current_pages and (
            current_chars + pg_chars > chunk_size
            or len(current_pages) >= max_pages
        ):
            text = pages_to_text(current_pages)
            chunks.append((current_pages, text))
            # Overlap: keep last page(s) up to `overlap` chars worth
            overlap_pages: list[PageText] = []
            overlap_total = 0
            for p in reversed(current_pages):
                if overlap_total + len(p.text) <= overlap:
                    overlap_pages.insert(0, p)
                    overlap_total += len(p.text)
                else:
                    break
            current_pages = overlap_pages
            current_chars = overlap_total

        current_pages.append(pg)
        current_chars += pg_chars

    if current_pages:
        text = pages_to_text(current_pages)
        chunks.append((current_pages, text))

    return chunks


def get_page_range(pages: list[PageText]) -> str:
    """Return a human-readable page range string, e.g. '12-15'."""
    nums = [pg.page_num for pg in pages]
    if not nums:
        return ""
    if len(nums) == 1:
        return str(nums[0])
    return f"{nums[0]}-{nums[-1]}"
