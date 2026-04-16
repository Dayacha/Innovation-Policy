"""
Gap filler for budget — closes gaps identified by gap_detector.

After gap_detector produces a reextract_queue, this module tries two passes
to find missing agency amounts before giving up:

  Phase 1 — Targeted DOCX text search (free, deterministic)
    Re-opens each flagged source document and searches ALL text
    (tables + paragraphs) for the missing agency's name variants.
    Uses broader matching than the main parser — no italic filter,
    no header exclusion — because we're looking for a specific agency.

  Phase 2 — LLM targeted extraction (cheap, ~$0.001/gap)
    For gaps still unresolved after Phase 1, extracts the relevant
    text sections from the DOCX and asks the LLM to find the amount.
    Only the sections mentioning the agency are sent (not the full doc),
    keeping cost very low.

After filling, the pipeline re-runs the canonical series + gap report
so the before/after improvement is visible.

Usage (standalone):
  python -m budget.gap_filler --country Australia

Called automatically from compile.py when --fill-gaps flag is set.
"""

from __future__ import annotations

import logging
import re
from pathlib import Path
from typing import Optional

import pandas as pd

from budget import config as cfg
from budget.canonical_series import _get_agencies_for_country

logger = logging.getLogger(__name__)

__all__ = ["fill_gaps"]

# ---------------------------------------------------------------------------
# Regex helpers
# ---------------------------------------------------------------------------

# Matches numbers like 1,234,567 or 1234567 or 1,234 — used to pull amounts
# from raw text. We strip commas before parsing.
_NUM_RE = re.compile(r"\b[\d]{1,3}(?:,\d{3})*(?:\.\d+)?\b|\b\d{4,}\b")

# Plausible amount range for a single R&D agency line in thousands AUD.
# Below $100k or above $50B → almost certainly a page number, date, or typo.
_AMOUNT_MIN = 100        # $100k
_AMOUNT_MAX = 50_000_000 # $50B


def _parse_amount(text: str) -> Optional[float]:
    """Extract the first plausible amount from a text string."""
    for m in _NUM_RE.finditer(text):
        raw = m.group().replace(",", "")
        try:
            val = float(raw)
            if _AMOUNT_MIN <= val <= _AMOUNT_MAX:
                return val
        except ValueError:
            continue
    return None


def _variants_lower(agency: dict) -> list[str]:
    canonical = agency["canonical_name"]
    variants = agency.get("name_variants", [canonical])
    return [v.lower() for v in variants]


# ---------------------------------------------------------------------------
# Phase 1 — Targeted DOCX text search
# ---------------------------------------------------------------------------

def _search_docx_tables(file_path: Path, agency: dict) -> list[dict]:
    """
    Re-parse ALL tables in a DOCX looking for the agency's name variants.
    No italic filter, no header exclusion — we want everything mentioning
    this specific agency.

    Returns list of dicts with entity_raw, amount_current, context.
    """
    try:
        import docx  # python-docx
    except ImportError:
        logger.error("python-docx not installed — cannot search DOCX tables")
        return []

    variants = _variants_lower(agency)
    results = []

    try:
        doc = docx.Document(str(file_path))
    except Exception as e:
        logger.warning(f"Cannot open {file_path.name}: {e}")
        return []

    for t_idx, table in enumerate(doc.tables):
        for r_idx, row in enumerate(table.rows):
            cells = [c.text.strip() for c in row.cells]
            # Deduplicate merged cells (python-docx repeats them)
            seen = set()
            unique_cells = []
            for c in cells:
                if c not in seen:
                    seen.add(c)
                    unique_cells.append(c)

            row_text = " ".join(unique_cells).lower()

            if not any(v in row_text for v in variants):
                continue

            # Found a row mentioning the agency — extract amounts
            entity_raw = unique_cells[0] if unique_cells else ""
            amounts = []
            for cell in unique_cells[1:]:  # skip entity column
                amt = _parse_amount(cell)
                if amt:
                    amounts.append(amt)

            if not amounts:
                # Maybe amount is on next row — check it
                continue

            # Take the largest amount (most likely the total, not sub-items)
            best_amount = max(amounts)
            results.append({
                "entity_raw": entity_raw,
                "amount_current": best_amount,
                "context": " | ".join(unique_cells),
                "table_index": t_idx,
                "row_index": r_idx,
                "method": "gap_fill_table",
            })

    return results


def _search_docx_paragraphs(file_path: Path, agency: dict) -> list[dict]:
    """
    Search paragraph text (non-table) for agency name variants + nearby amounts.
    Useful when an agency's amount is in a narrative or summary section.
    """
    try:
        import docx
    except ImportError:
        return []

    variants = _variants_lower(agency)
    results = []

    try:
        doc = docx.Document(str(file_path))
    except Exception as e:
        return []

    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]

    for i, para in enumerate(paragraphs):
        para_lower = para.lower()
        if not any(v in para_lower for v in variants):
            continue

        # Found a paragraph mentioning the agency — look for amounts
        # in this paragraph and the next 3
        context_paras = paragraphs[i:min(i + 4, len(paragraphs))]
        context = " ".join(context_paras)
        amt = _parse_amount(context)
        if amt:
            results.append({
                "entity_raw": para[:120],
                "amount_current": amt,
                "context": context[:300],
                "table_index": -1,
                "row_index": i,
                "method": "gap_fill_paragraph",
            })
            break  # one match per paragraph block is enough

    return results


# ---------------------------------------------------------------------------
# Phase 2 — LLM targeted extraction
# ---------------------------------------------------------------------------

_LLM_SYSTEM = """You are a government budget analyst. You will be given a section
of an Australian Finance Bill (Appropriations Act). Your task is to find the
budget appropriation for a specific agency.

Respond ONLY with valid JSON:
{
  "found": true or false,
  "amount": <number in the document's unit, or null if not found>,
  "unit": "thousands" or "dollars" or "unknown",
  "raw_text": "<the exact line you found, or empty string>",
  "confidence": <0.0 to 1.0>
}

If the agency appears with multiple amounts, return the largest (it is likely
the total appropriation rather than a sub-component).
"""


def _llm_extract_from_docx(
    file_path: Path,
    agency: dict,
    config: dict,
    year: int,
) -> Optional[dict]:
    """
    Extract text sections from a DOCX that mention the agency,
    then ask the LLM to identify the budget amount.

    Returns dict with amount_current and confidence, or None.
    """
    try:
        import docx
    except ImportError:
        return None

    try:
        from budget.llm_client import BudgetLLMClient
    except ImportError:
        logger.error("LLM client not available")
        return None

    variants = _variants_lower(agency)
    canonical = agency["canonical_name"]

    # Collect all text snippets mentioning the agency
    try:
        doc = docx.Document(str(file_path))
    except Exception as e:
        logger.warning(f"Cannot open {file_path.name}: {e}")
        return None

    snippets = []

    # From tables
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            row_text = " ".join(cells).lower()
            if any(v in row_text for v in variants):
                snippets.append(" | ".join(c for c in cells if c))

    # From paragraphs
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    for i, para in enumerate(paragraphs):
        if any(v in para.lower() for v in variants):
            # Include surrounding context
            ctx = " ".join(paragraphs[max(0, i-1):min(i+3, len(paragraphs))])
            snippets.append(ctx)

    if not snippets:
        logger.debug(f"No text found for {canonical} in {file_path.name}")
        return None

    # Limit context to first 5 snippets to keep prompt small
    context_text = "\n\n".join(snippets[:5])
    if len(context_text) > 4000:
        context_text = context_text[:4000]

    user_prompt = (
        f"Document: {file_path.name} (year {year})\n"
        f"Find the budget appropriation for: {canonical}\n"
        f"Also known as: {', '.join(agency.get('name_variants', [])[:5])}\n\n"
        f"Relevant document sections:\n{context_text}"
    )

    try:
        client = BudgetLLMClient(config)
        result = client.call_json(
            system_prompt=_LLM_SYSTEM,
            user_prompt=user_prompt,
            max_tokens=256,
            operation=client.OP_EXTRACT,
        )
    except Exception as e:
        logger.warning(f"LLM call failed for {canonical} in {file_path.name}: {e}")
        return None

    if "_parse_error" in result:
        logger.warning(f"LLM JSON parse error for {canonical}: {result['_parse_error']}")
        return None

    if not result.get("found"):
        return None

    amount = result.get("amount")
    if not amount:
        return None

    try:
        amount = float(amount)
    except (TypeError, ValueError):
        return None

    if not (_AMOUNT_MIN <= amount <= _AMOUNT_MAX):
        return None

    return {
        "amount_current": amount,
        "entity_raw": f"{canonical} (LLM extracted)",
        "context": result.get("raw_text", ""),
        "confidence": float(result.get("confidence", 0.7)),
        "method": "gap_fill_llm",
    }


# ---------------------------------------------------------------------------
# Main entry point
# ---------------------------------------------------------------------------

def fill_gaps(
    gap_df: pd.DataFrame,
    country: str,
    config: dict,
    pdf_root: Path = cfg.PDF_ROOT,
    use_llm: bool = True,
) -> pd.DataFrame:
    """
    For each gap in gap_df with action='reextract', attempt to find the missing
    amount by re-searching the source documents.

    Phase 1: Targeted DOCX text search (free)
    Phase 2: LLM targeted extraction (cheap) — only if use_llm=True and Phase 1 failed

    Returns a DataFrame of newly found rows in raw_rows format, ready to be
    appended to raw_df and fed back into build_canonical_series().

    The gap_df is also updated in-place with the results (action changed to
    'filled' or 'fill_failed').
    """
    reextract = gap_df[gap_df["action"] == "reextract"].copy()
    if reextract.empty:
        logger.info(f"[{country}] No gaps to fill")
        return pd.DataFrame()

    agencies = {a["canonical_name"]: a for a in _get_agencies_for_country(country)}
    country_dir = pdf_root / country

    new_rows = []
    filled = 0
    failed = 0

    logger.info(
        f"[{country}] Gap filler: {len(reextract)} gaps to process "
        f"(Phase 1: text search, Phase 2: {'LLM' if use_llm else 'disabled'})"
    )

    # Group by (year, source_file) so we open each file once per year
    for (year, canonical), group in reextract.groupby(["year", "canonical_name"]):
        agency = agencies.get(canonical)
        if not agency:
            logger.debug(f"Agency {canonical} not in registry — skipping gap fill")
            continue

        # Find the source file for this year
        source_file = _find_source_file(country_dir, year)
        if not source_file:
            logger.debug(f"No source file found for {country} {year}")
            failed += len(group)
            continue

        # ── Phase 1: table + paragraph search ────────────────────────────────
        table_hits = _search_docx_tables(source_file, agency)
        para_hits = _search_docx_paragraphs(source_file, agency)
        all_hits = table_hits + para_hits

        found_row = None
        if all_hits:
            # Pick the hit with the largest amount (most likely the agency total)
            best = max(all_hits, key=lambda h: h["amount_current"])
            found_row = best
            logger.info(
                f"[{country}] Phase 1 found {canonical} {year}: "
                f"{best['amount_current']:,.0f} in {source_file.name} "
                f"via {best['method']}"
            )

        # ── Phase 2: LLM extraction ───────────────────────────────────────────
        if found_row is None and use_llm:
            llm_result = _llm_extract_from_docx(source_file, agency, config, year)
            if llm_result:
                found_row = llm_result
                logger.info(
                    f"[{country}] Phase 2 found {canonical} {year}: "
                    f"{llm_result['amount_current']:,.0f} in {source_file.name} "
                    f"(confidence={llm_result['confidence']:.2f})"
                )

        if found_row:
            new_rows.append({
                "country": country,
                "year": year,
                "source_file": source_file.name,
                "table_index": found_row.get("table_index", -1),
                "row_index": found_row.get("row_index", -1),
                "section_name": canonical,
                "entity_raw": found_row["entity_raw"],
                "amount_current": found_row["amount_current"],
                "amount_prior": None,
                "is_header_row": False,
                "is_total_row": False,
                "has_italic_entity": False,
                "cells_raw": found_row.get("context", ""),
                "canonical_name": canonical,
                "unit_note": f"gap_fill ({found_row['method']})",
            })
            # Update gap_df
            gap_df.loc[group.index, "action"] = "filled"
            gap_df.loc[group.index, "diagnosis"] = (
                f"Gap filled via {found_row['method']}: "
                f"{found_row['amount_current']:,.0f}"
            )
            filled += 1
        else:
            gap_df.loc[group.index, "action"] = "fill_failed"
            failed += 1

    logger.info(
        f"[{country}] Gap filler complete: {filled} filled, {failed} still missing"
    )

    return pd.DataFrame(new_rows) if new_rows else pd.DataFrame()


def _find_source_file(country_dir: Path, year: int) -> Optional[Path]:
    """
    Find the primary (lowest Act number) DOCX for a given country/year.
    Prefers 'No1' files as the primary appropriation act.
    """
    if not country_dir.exists():
        return None

    _YEAR_PAT = re.compile(r"(?<![0-9])(" + str(year) + r")(?![0-9])")
    _ACT_NO = re.compile(r"\bNo\.?\s*(\d+)\b", re.IGNORECASE)

    candidates = []
    for path in country_dir.iterdir():
        if not path.suffix.lower() in (".docx",):
            continue
        if _YEAR_PAT.search(path.stem):
            m = _ACT_NO.search(path.stem)
            act_num = int(m.group(1)) if m else 999
            candidates.append((act_num, path))

    if not candidates:
        return None

    # Return the primary act (lowest number = annual appropriation, not supplementary)
    candidates.sort(key=lambda x: x[0])
    return candidates[0][1]


# ---------------------------------------------------------------------------
# CLI
# ---------------------------------------------------------------------------

if __name__ == "__main__":
    import argparse
    import yaml

    logging.basicConfig(level=logging.INFO, format="%(levelname)s %(name)s — %(message)s")

    parser = argparse.ArgumentParser(description="Fill gaps in canonical R&D series")
    parser.add_argument("--country", required=True)
    parser.add_argument("--gap-report", help="Path to gap_report CSV (default: auto)")
    parser.add_argument("--no-llm", action="store_true", help="Skip LLM phase")
    parser.add_argument("--config", default="config.yaml")
    args = parser.parse_args()

    with open(args.config) as f:
        config = yaml.safe_load(f)

    cname = args.country.lower().replace(" ", "_")
    gap_path = Path(args.gap_report) if args.gap_report else (
        cfg.OUTPUT_DIR / f"{cname}_gap_report.csv"
    )

    if not gap_path.exists():
        print(f"Gap report not found: {gap_path}")
        print("Run compile first: python -m budget.compile --country ...")
        exit(1)

    gap_df = pd.read_csv(gap_path)
    new_rows_df = fill_gaps(
        gap_df=gap_df,
        country=args.country,
        config=config,
        use_llm=not args.no_llm,
    )

    # Save updated gap report
    gap_df.to_csv(gap_path, index=False)

    n_reextract = len(gap_df[gap_df["action"] == "reextract"])
    n_filled = len(gap_df[gap_df["action"] == "filled"])
    n_failed = len(gap_df[gap_df["action"] == "fill_failed"])

    print(f"\n=== Gap filler results for {args.country} ===")
    print(f"Filled:        {n_filled}")
    print(f"Still missing: {n_failed}")
    print(f"Untouched:     {n_reextract}")

    if not new_rows_df.empty:
        print(f"\nNew rows found ({len(new_rows_df)}):")
        print(new_rows_df[["year", "canonical_name", "amount_current", "unit_note"]].to_string())
        print(f"\nRe-run compile to incorporate these into the series:")
        print(f"  python -m budget.compile --country {args.country} --config {args.config} --no-entity-dedup")
