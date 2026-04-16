"""
Deterministic DOCX table parser for budget.

Extracts ALL (entity, amount) pairs from every table in a .docx file
using python-docx structural metadata. No LLM. No filtering.

Key design decisions:
- Current year = plain text (not italic)
- Prior year  = italic text → always tagged and EXCLUDED from output
- Every row in every table is extracted — nothing is filtered here
- Section headers are tracked from bold / all-caps rows
- Output is a flat list of RawRow dicts, one per extracted line

This works for any .docx regardless of country or language because
italic/plain is a structural XML property, not a content inference.

After this runs, the LLM is only needed for:
  1. Classifying entity names (once per unique name, cached forever)
  2. Non-table documents (paragraphs, PDFs) — fallback to existing extractor
"""

from __future__ import annotations

import logging
import re
from dataclasses import dataclass, asdict, field
from pathlib import Path
from typing import Optional

logger = logging.getLogger(__name__)

__all__ = ["parse_docx", "RawRow"]


# ---------------------------------------------------------------------------
# Data model
# ---------------------------------------------------------------------------

@dataclass
class RawRow:
    """One extracted row from a DOCX table. No filtering applied."""

    source_file: str = ""
    country: str = ""
    year: int = 0
    page_number: int = 0        # approximate — table index used if page not available
    table_index: int = 0
    row_index: int = 0

    # Section context (tracked from bold/header rows above this row)
    section_name: str = ""      # as printed (may be non-English)

    # Entity: the text in the first non-empty cell of the row
    entity_raw: str = ""        # exactly as printed, original language

    # Amounts: we capture all numeric cells found in the row
    # amount_current: the plain-text (non-italic) numeric value
    # amount_prior:   the italic numeric value (prior year — EXCLUDED from series)
    amount_current: Optional[float] = None
    amount_prior: Optional[float] = None

    # Raw cells as text (for debugging / audit)
    cells_raw: list = field(default_factory=list)

    # Flags
    is_header_row: bool = False     # bold or all-caps row, likely a section header
    is_total_row: bool = False      # "Total:" prefix in entity text
    has_italic_entity: bool = False # entity cell itself is italic (rare)

    def to_dict(self) -> dict:
        d = asdict(self)
        d["cells_raw"] = " | ".join(str(c) for c in d["cells_raw"])
        return d


# ---------------------------------------------------------------------------
# Parsing helpers
# ---------------------------------------------------------------------------

_RE_NUMBER = re.compile(r"^-?[\d,]+\.?\d*$")
_RE_TOTAL = re.compile(r"^(total|totals?:|sum)\b", re.IGNORECASE)


def _parse_number(text: str) -> Optional[float]:
    """Parse a cell value like '1,234,567' or '-' into a float. Returns None for blanks/dashes."""
    text = text.strip().replace(",", "").replace(" ", "")
    if not text or text in ("-", "—", "–", "*", "nil", "n/a"):
        return None
    try:
        return float(text)
    except ValueError:
        return None


def _cell_italic(cell) -> bool:
    """Return True if ALL non-empty runs in a cell are italic."""
    runs_with_text = [
        run
        for para in cell.paragraphs
        for run in para.runs
        if run.text.strip()
    ]
    if not runs_with_text:
        return False
    return all(run.italic for run in runs_with_text)


def _cell_bold(cell) -> bool:
    """Return True if ANY run in a cell is bold."""
    return any(
        run.bold
        for para in cell.paragraphs
        for run in para.runs
        if run.text.strip()
    )


def _cell_all_caps(text: str) -> bool:
    """Return True if text is all-uppercase (likely a section header)."""
    letters = [c for c in text if c.isalpha()]
    return len(letters) >= 4 and all(c.isupper() for c in letters)


def _is_header_row(row_cells: list[dict]) -> bool:
    """
    Heuristic: a row is a section header if:
    - It has only one non-empty cell, or
    - The first cell is bold or ALL CAPS, or
    - All non-empty cells contain no numbers
    """
    texts = [c["text"] for c in row_cells if c["text"]]
    if not texts:
        return False

    # Only one cell has content → likely a section heading
    if len(texts) == 1:
        t = texts[0]
        return not _parse_number(t) and len(t) > 3

    # First cell is bold or ALL CAPS and no cell is numeric
    first = row_cells[0]["text"]
    if (_cell_all_caps(first) or row_cells[0].get("bold")) and not any(
        _parse_number(c["text"]) for c in row_cells
    ):
        return True

    return False


def _extract_amounts(row_cells: list[dict]) -> tuple[Optional[float], Optional[float]]:
    """
    Extract current-year (plain) and prior-year (italic) amounts from a row.

    Australian budget tables typically have multiple amount columns
    (Departmental, Administered, Total). We take the LAST non-null numeric
    value as the primary amount — this is usually the Total column.
    Italic cells → prior year; plain cells → current year.
    """
    current_values = []
    prior_values = []

    for cell in row_cells[1:]:  # skip first cell (entity name)
        val = _parse_number(cell["text"])
        if val is None:
            continue
        if cell["italic"]:
            prior_values.append(val)
        else:
            current_values.append(val)

    # Take the last value (usually "Total" column)
    current = current_values[-1] if current_values else None
    prior = prior_values[-1] if prior_values else None

    return current, prior


# ---------------------------------------------------------------------------
# Main parser
# ---------------------------------------------------------------------------

def parse_docx(
    path: Path,
    source_file: str = "",
    country: str = "",
    year: int = 0,
) -> list[RawRow]:
    """
    Parse all tables in a .docx file and return flat list of RawRow objects.

    Parameters
    ----------
    path        : Path to the .docx file
    source_file : Original filename (for audit trail)
    country     : Country name (passed through to rows)
    year        : Budget year (passed through to rows)

    Returns
    -------
    List of RawRow objects. One row per table row that has at least one
    non-empty cell. Prior-year rows are included but tagged amount_prior only.
    Completely empty rows are skipped.
    """
    try:
        import docx as python_docx
    except ImportError:
        logger.error("python-docx not installed: pip install python-docx")
        return []

    try:
        doc = python_docx.Document(path)
    except Exception as e:
        logger.warning(f"Cannot open {path}: {e}")
        return []

    rows: list[RawRow] = []
    current_section = ""

    for t_idx, table in enumerate(doc.tables):
        for r_idx, row in enumerate(table.rows):

            # Build cell metadata
            cells = []
            for cell in row.cells:
                text = cell.text.strip()
                # Deduplicate merged cells (python-docx repeats merged cells)
                if cells and cells[-1]["text"] == text and text == "":
                    continue
                cells.append({
                    "text": text,
                    "italic": _cell_italic(cell),
                    "bold": _cell_bold(cell),
                })

            # Skip completely empty rows
            if not any(c["text"] for c in cells):
                continue

            # Detect header rows — update section context
            if _is_header_row(cells):
                header_text = " ".join(c["text"] for c in cells if c["text"])
                # Only update section if it looks like a real section name (not a column header)
                if not any(
                    word in header_text.lower()
                    for word in ["departmental", "administered", "total", "$'000", "appropriation"]
                ):
                    current_section = header_text
                rows.append(RawRow(
                    source_file=source_file,
                    country=country,
                    year=year,
                    table_index=t_idx,
                    row_index=r_idx,
                    section_name=current_section,
                    entity_raw=header_text,
                    is_header_row=True,
                    cells_raw=[c["text"] for c in cells],
                ))
                continue

            # Entity name: first non-empty cell
            entity_text = next((c["text"] for c in cells if c["text"]), "")
            entity_italic = cells[0]["italic"] if cells else False

            # Extract amounts
            amount_current, amount_prior = _extract_amounts(cells)

            # Skip rows with no entity and no amounts
            if not entity_text and amount_current is None and amount_prior is None:
                continue

            # Prior-year-only rows: blank entity + only italic amounts
            # Still record them but only in amount_prior
            if not entity_text and amount_prior is not None and amount_current is None:
                # This is a pure prior-year row — attach to previous entity in audit
                rows.append(RawRow(
                    source_file=source_file,
                    country=country,
                    year=year,
                    table_index=t_idx,
                    row_index=r_idx,
                    section_name=current_section,
                    entity_raw="",   # prior year continuation
                    amount_current=None,
                    amount_prior=amount_prior,
                    has_italic_entity=True,
                    cells_raw=[c["text"] for c in cells],
                ))
                continue

            is_total = bool(_RE_TOTAL.match(entity_text)) or "total" in entity_text.lower()

            rows.append(RawRow(
                source_file=source_file,
                country=country,
                year=year,
                table_index=t_idx,
                row_index=r_idx,
                section_name=current_section,
                entity_raw=entity_text,
                amount_current=amount_current,
                amount_prior=amount_prior,
                is_total_row=is_total,
                has_italic_entity=entity_italic,
                cells_raw=[c["text"] for c in cells],
            ))

    logger.info(
        f"Parsed {path.name}: {len(doc.tables)} tables, "
        f"{len(rows)} rows extracted "
        f"({sum(1 for r in rows if r.amount_current is not None)} with current-year amounts)"
    )
    return rows


# ---------------------------------------------------------------------------
# Batch helper: parse all DOCX files for a country/year range
# ---------------------------------------------------------------------------

def parse_country_docx_files(
    pdf_root: Path,
    country: str,
    year_range: Optional[tuple[int, int]] = None,
) -> list[RawRow]:
    """
    Parse all DOCX files for a country and return combined RawRow list.

    Parameters
    ----------
    pdf_root   : Root directory containing country subdirectories
    country    : Country directory name
    year_range : (start, end) inclusive. None = all years.
    """
    import re as _re

    _YEAR_PAT = _re.compile(r"(?<![0-9])(1[89]\d{2}|20[012]\d)(?![0-9])")

    country_dir = pdf_root / country
    if not country_dir.exists():
        logger.warning(f"Country directory not found: {country_dir}")
        return []

    all_rows: list[RawRow] = []

    docx_files = sorted(country_dir.glob("*.docx")) + sorted(country_dir.glob("*.DOCX"))
    for path in docx_files:
        m = _YEAR_PAT.search(path.stem)
        if not m:
            continue
        year = int(m.group(1))
        if year_range and not (year_range[0] <= year <= year_range[1]):
            continue

        file_rows = parse_docx(
            path,
            source_file=path.name,
            country=country,
            year=year,
        )
        all_rows.extend(file_rows)

    logger.info(
        f"[{country}] Parsed {len(docx_files)} DOCX files → "
        f"{len(all_rows)} total rows, "
        f"{sum(1 for r in all_rows if r.amount_current is not None)} with current-year amounts"
    )
    return all_rows


# ---------------------------------------------------------------------------
# CLI — quick test
# ---------------------------------------------------------------------------

if __name__ == "__main__":
    import argparse
    import csv
    import sys

    logging.basicConfig(level=logging.INFO, format="%(levelname)s — %(message)s")

    parser = argparse.ArgumentParser(description="Parse DOCX tables deterministically")
    parser.add_argument("path", help="Path to .docx file or country directory")
    parser.add_argument("--country", default="Australia")
    parser.add_argument("--year", type=int, default=0)
    parser.add_argument("--output", default="-", help="Output CSV path (- for stdout)")
    args = parser.parse_args()

    p = Path(args.path)
    if p.is_file():
        results = parse_docx(p, source_file=p.name, country=args.country, year=args.year)
    else:
        results = parse_country_docx_files(p.parent, p.name)

    # Filter to rows with current-year amounts only
    data_rows = [r for r in results if r.amount_current is not None]
    print(f"\nRows with current-year amounts: {len(data_rows)}", file=sys.stderr)

    if args.output == "-":
        writer = csv.DictWriter(sys.stdout, fieldnames=list(RawRow.__dataclass_fields__.keys()))
        writer.writeheader()
        for r in data_rows:
            writer.writerow(r.to_dict())
    else:
        import pandas as pd
        df = pd.DataFrame([r.to_dict() for r in data_rows])
        df.to_csv(args.output, index=False)
        print(f"Saved to {args.output}", file=sys.stderr)
