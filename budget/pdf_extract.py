"""PDF and DOCX text extraction stage with OCR fallback when direct extraction is weak."""

from pathlib import Path
from typing import Optional

import fitz
import pandas as pd

from budget.config import MIN_ALNUM_RATIO, MIN_DIRECT_TEXT_CHARS, PAGE_EXTRACTION_FILE
from budget.ocr_utils import is_ocr_available, run_ocr_on_page
from budget.utils import logger, text_quality_metrics


def _extract_docx_pages(filepath: Path) -> list[dict]:
    """Extract text from a .docx file, returning one 'page' per table plus one for paragraphs.

    Tables are converted to tab-separated rows so that structured data
    (like Australian Appropriation Act schedule tables) remains parseable.
    """
    try:
        from docx import Document  # python-docx
    except ImportError:
        logger.warning("python-docx not installed; skipping docx extraction for %s", filepath.name)
        return []

    try:
        doc = Document(filepath)
    except Exception as exc:
        logger.error("Failed to open docx %s: %s", filepath, exc)
        return []

    pages = []

    # Page 1: all paragraph text
    para_text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
    if para_text.strip():
        pages.append({"page_number": 1, "text": para_text, "extraction_method": "docx_paragraphs"})

    # Pages 2+: one page per table (tab-separated cells)
    for i, table in enumerate(doc.tables):
        rows_text = []
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            # Deduplicate merged cells (python-docx repeats merged cells)
            seen = []
            prev = None
            for c in cells:
                if c != prev:
                    seen.append(c)
                prev = c
            row_line = "\t".join(seen)
            if row_line.strip():
                rows_text.append(row_line)
        table_text = "\n".join(rows_text)
        if table_text.strip():
            pages.append({
                "page_number": i + 2,
                "text": table_text,
                "extraction_method": "docx_table",
            })

    return pages


def _needs_ocr(text: str) -> bool:
    """Heuristic for deciding when direct extraction quality is too low."""
    char_count, alnum_ratio = text_quality_metrics(text)
    if char_count < MIN_DIRECT_TEXT_CHARS:
        return True
    if alnum_ratio < MIN_ALNUM_RATIO:
        return True
    return False


def extract_text_for_inventory(inventory_df: pd.DataFrame) -> tuple[pd.DataFrame, pd.DataFrame]:
    """Extract page-level text for all files in inventory with simple caching.

    If page_text.csv already contains rows for a file with the same content hash,
    those rows are reused instead of re-extracting, even when the filename changed.
    """
    page_records = []
    summary_records = []
    ocr_ready = is_ocr_available()

    cache_df = None
    if PAGE_EXTRACTION_FILE.exists():
        try:
            cache_df = pd.read_csv(PAGE_EXTRACTION_FILE)
        except Exception:
            cache_df = None

    if not ocr_ready:
        logger.warning(
            "OCR is not available (missing pytesseract or Tesseract binary). "
            "Scanned pages may stay empty."
        )

    for row in inventory_df.itertuples(index=False):
        file_id = row.file_id
        filepath = Path(row.filepath)
        country_guess = row.country_guess
        year_guess = row.year_guess
        file_size = getattr(row, "file_size", None)
        content_hash = getattr(row, "content_hash", None)

        # Reuse cached pages if available and the PDF bytes match, even if renamed.
        if cache_df is not None:
            cached_rows = pd.DataFrame()
            if content_hash and "content_hash" in cache_df.columns:
                cached_rows = cache_df[cache_df["content_hash"] == content_hash]
            elif "file_size" in cache_df.columns:
                cached_rows = cache_df[cache_df["filepath"] == str(filepath)]
                if file_size is not None:
                    cached_rows = cached_rows[cached_rows.get("file_size") == file_size]
            if not cached_rows.empty:
                cached_rows = cached_rows.copy()
                # Deduplicate by page_number to prevent exponential CSV growth
                # when the cache itself already contains duplicate rows
                cached_rows = cached_rows.drop_duplicates(subset=["page_number"], keep="last")
                cached_rows["file_id"] = file_id
                cached_rows["filepath"] = str(filepath)
                cached_rows["file_size"] = file_size
                cached_rows["country_guess"] = country_guess
                cached_rows["year_guess"] = year_guess
                cached_rows["content_hash"] = content_hash
                page_records.extend(cached_rows.to_dict(orient="records"))
                summary_records.append(
                    {
                        "file_id": file_id,
                        "filepath": str(filepath),
                        "country_guess": country_guess,
                        "year_guess": year_guess,
                        "content_hash": content_hash,
                        "file_size": file_size,
                        "total_pages": cached_rows["page_number"].max(),
                        "direct_pages": (cached_rows["extraction_method"] == "direct_text").sum(),
                        "ocr_pages": (cached_rows["extraction_method"] == "ocr_fallback").sum(),
                        "error_pages": (cached_rows["extraction_method"].str.contains("error").sum() if "extraction_method" in cached_rows else 0),
                        "status": "cached_by_content" if content_hash else "cached",
                    }
                )
                logger.info("Reused cached text for %s", filepath.name)
                continue

        logger.info("Extracting file: %s", filepath.name)
        file_total_pages = 0
        direct_pages = 0
        ocr_pages = 0
        error_pages = 0
        file_status = "ok"

        # ── DOCX branch ──────────────────────────────────────────────────────
        if filepath.suffix.lower() == ".docx":
            try:
                docx_pages = _extract_docx_pages(filepath)
                file_total_pages = len(docx_pages)
                direct_pages = file_total_pages
                for pg in docx_pages:
                    char_count = len(pg["text"])
                    page_records.append({
                        "file_id": file_id,
                        "filepath": str(filepath),
                        "file_size": file_size,
                        "content_hash": content_hash,
                        "country_guess": country_guess,
                        "year_guess": year_guess,
                        "page_number": pg["page_number"],
                        "extraction_method": pg["extraction_method"],
                        "text": pg["text"],
                        "char_count": char_count,
                    })
            except Exception as exc:
                file_status = f"error: {exc}"
                logger.error("Failed to process docx %s: %s", filepath, exc)
                error_pages += 1
            summary_records.append({
                "file_id": file_id,
                "filepath": str(filepath),
                "file_size": file_size,
                "content_hash": content_hash,
                "country_guess": country_guess,
                "year_guess": year_guess,
                "total_pages": file_total_pages,
                "direct_pages": direct_pages,
                "ocr_pages": 0,
                "error_pages": error_pages,
                "status": file_status,
            })
            continue
        # ── PDF branch ───────────────────────────────────────────────────────

        try:
            with fitz.open(filepath) as doc:
                file_total_pages = len(doc)
                for page_idx, page in enumerate(doc):
                    page_number = page_idx + 1
                    extraction_method = "direct_text"
                    text = page.get_text("text") or ""

                    if _needs_ocr(text):
                        if ocr_ready:
                            ocr_text = run_ocr_on_page(page)
                            if ocr_text.strip():
                                text = ocr_text
                                extraction_method = "ocr_fallback"
                                ocr_pages += 1
                            else:
                                extraction_method = "direct_text_low_quality"
                                direct_pages += 1
                        else:
                            extraction_method = "direct_text_no_ocr"
                            direct_pages += 1
                    else:
                        direct_pages += 1

                    char_count = len(text)
                    page_records.append(
                        {
                            "file_id": file_id,
                            "filepath": str(filepath),
                            "file_size": file_size,
                            "content_hash": content_hash,
                            "country_guess": country_guess,
                            "year_guess": year_guess,
                            "page_number": page_number,
                            "extraction_method": extraction_method,
                            "text": text,
                            "char_count": char_count,
                        }
                    )
        except Exception as exc:
            file_status = f"error: {exc}"
            logger.error("Failed to process file %s: %s", filepath, exc)
            error_pages += 1

        summary_records.append(
            {
                "file_id": file_id,
                "filepath": str(filepath),
                "file_size": file_size,
                "content_hash": content_hash,
                "country_guess": country_guess,
                "year_guess": year_guess,
                "total_pages": file_total_pages,
                "direct_pages": direct_pages,
                "ocr_pages": ocr_pages,
                "error_pages": error_pages,
                "status": file_status,
            }
        )

    # ── Retain cached text for PDFs no longer on disk ────────────────────────
    # This allows deleting PDFs after extraction while keeping results intact.
    # Any file in page_text.csv that isn't in the current inventory is an
    # "orphaned" cached entry — include it so downstream stages still see it.
    if cache_df is not None:
        current_paths = set(str(Path(row.filepath)) for row in inventory_df.itertuples(index=False))
        orphaned = cache_df[~cache_df["filepath"].isin(current_paths)].copy()
        if "content_hash" in cache_df.columns and "content_hash" in inventory_df.columns:
            current_hashes = set(
                str(getattr(row, "content_hash"))
                for row in inventory_df.itertuples(index=False)
                if getattr(row, "content_hash", None)
            )
            if current_hashes:
                orphaned = orphaned[~orphaned["content_hash"].astype(str).isin(current_hashes)]
        if not orphaned.empty:
            logger.info(
                "Retaining cached text for %d previously processed PDF(s) no longer on disk (%d pages)",
                orphaned["filepath"].nunique(),
                len(orphaned),
            )
            page_records.extend(orphaned.to_dict(orient="records"))

    pages_df = pd.DataFrame(page_records)
    summary_df = pd.DataFrame(summary_records)
    return pages_df, summary_df
